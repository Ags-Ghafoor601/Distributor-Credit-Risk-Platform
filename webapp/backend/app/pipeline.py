"""
Shared Scoring Pipeline Module
Problem 01: Distributor Credit Risk on Gut Feel

This module contains the SAME logic as score_dealers.py and robust_ingestion.py,
refactored into importable functions operating on in-memory DataFrames instead
of fixed file paths -- necessary because the web app accepts uploaded files
rather than reading from disk.

CRITICAL: every function here was verified to reproduce EXACT known values
from the validated CLI pipeline (D0080 = 418, 220 dealers, 178/36/6 split)
before this module was considered correct. See verify_against_cli.py.
"""

import io
from datetime import date
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CUTOFF_DATE = pd.Timestamp("2025-01-01")
ANNUAL_PKR_EROSION = 0.18
MIN_INVOICES = 3

REASON_LABELS = {
    "payment_delay_severity": "Late and inconsistent payment timing",
    "bounce_rate_lifetime": "Cheque bounce history",
    "real_exposure_pkr": "Inflation-adjusted credit exposure",
    "order_frequency_trend": "Declining order activity",
    "salesman_default_rate_loo": "Salesman's track record with similar dealers",
    "territory_default_rate_loo": "Elevated risk in dealer's territory",
}

TRUE_SET = {"y", "yes", "1", "true", "bounced"}
FALSE_SET = {"n", "no", "0", "false", "", "ok", "nan"}

# EXACT seasonal windows from robust_ingestion.py -- must stay identical to
# the validated CLI pipeline. If these ever need to change, change them
# in BOTH places, or better, extract to a shared constants file.
SEASONAL_WINDOWS = [
    (date(2023, 3, 10), date(2023, 5, 5)), (date(2023, 6, 15), date(2023, 7, 10)),
    (date(2024, 2, 28), date(2024, 4, 25)), (date(2024, 6, 5), date(2024, 6, 30)),
    (date(2025, 2, 15), date(2025, 4, 15)), (date(2025, 5, 25), date(2025, 6, 20)),
]


def is_seasonal(d):
    if pd.isna(d):
        return False
    d = d.date() if hasattr(d, "date") else d
    return any(s <= d <= e for s, e in SEASONAL_WINDOWS)


# ---------------------------------------------------------------------------
# INGESTION (ported from robust_ingestion.py)
# ---------------------------------------------------------------------------
def parse_messy_date(val):
    if pd.isna(val) or str(val).strip() in ("", "N/A", "NA", "-", "nil", "pending"):
        return pd.NaT
    s = str(val).strip()
    if s.isdigit() and 40000 < int(s) < 48000:
        try:
            return pd.Timestamp("1899-12-30") + pd.Timedelta(days=int(s))
        except Exception:
            return pd.NaT
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m-%d-%Y", "%d-%b-%y", "%Y.%m.%d", "%d-%b-%Y"):
        try:
            return pd.to_datetime(s, format=fmt)
        except Exception:
            continue
    try:
        return pd.to_datetime(s, dayfirst=True, errors="raise")
    except Exception:
        return pd.NaT


def parse_messy_amount(val):
    if pd.isna(val):
        return np.nan
    s = str(val).strip()
    if s in ("", "N/A", "NA", "-", "nil"):
        return np.nan
    s = s.strip("()").replace("Rs.", "").replace("Rs", "").replace("PKR", "")
    s = s.replace(",", "").replace("/-", "").strip()
    try:
        return float(s)
    except ValueError:
        return np.nan


def parse_bool(val):
    s = str(val).strip().lower()
    if s in TRUE_SET:
        return True
    if s in FALSE_SET:
        return False
    return np.nan


def clean_transactions(raw_bytes: bytes, valid_dealer_ids: set) -> tuple[pd.DataFrame, dict]:
    """
    Cleans a raw (potentially messy) uploaded transactions CSV.
    Returns (cleaned_dataframe, quality_report_dict).
    Same rules as robust_ingestion.py -- NaN detection uses .isna() explicitly
    (not string-sentinel matching), per the bug fix caught during the
    Day 2 stress test.
    """
    raw = pd.read_csv(io.BytesIO(raw_bytes), dtype=str)
    raw.columns = [c.strip() for c in raw.columns]

    COLUMN_MAP = {
        "Txn ID": "transaction_id", "Dealer Code": "dealer_id", "Invoice Dt": "invoice_date",
        "Due Date": "due_date", "Paid On": "payment_date", "Amount (Rs)": "amount_pkr",
        "Mode": "payment_method", "Bounced?": "cheque_bounced",
        "transaction_id": "transaction_id", "dealer_id": "dealer_id",
        "invoice_date": "invoice_date", "due_date": "due_date", "payment_date": "payment_date",
        "amount_pkr": "amount_pkr", "payment_method": "payment_method", "cheque_bounced": "cheque_bounced",
    }
    df = raw.rename(columns=COLUMN_MAP)
    df = df[[c for c in df.columns if not c.startswith("Unnamed") and c != "Notes"]]

    report = {"input_rows": len(df)}

    df = df.drop_duplicates()

    missing_dealer_raw = df["dealer_id"].isna() if "dealer_id" in df.columns else pd.Series([False] * len(df))
    df["dealer_id"] = df["dealer_id"].astype(str).str.strip()
    missing_dealer = missing_dealer_raw | df["dealer_id"].isin(["nan", "", "None", "N/A", "NA"])
    n_missing = int(missing_dealer.sum())
    df = df[~missing_dealer]

    orphan_mask = ~df["dealer_id"].isin(valid_dealer_ids)
    n_orphan = int(orphan_mask.sum())
    df = df[~orphan_mask]

    for col in ["invoice_date", "due_date", "payment_date"]:
        df[col] = df[col].apply(parse_messy_date)
    before = len(df)
    df = df[df["invoice_date"].notna() & df["due_date"].notna()]
    n_bad_dates = before - len(df)

    df["amount_pkr"] = df["amount_pkr"].apply(parse_messy_amount)
    before = len(df)
    df = df[df["amount_pkr"].notna() & (df["amount_pkr"] > 0)]
    n_bad_amount = before - len(df)

    df["cheque_bounced"] = df["cheque_bounced"].apply(parse_bool)
    before = len(df)
    df = df[df["cheque_bounced"].notna()]
    n_bad_bool = before - len(df)

    before = len(df)
    df = df[df["due_date"] >= df["invoice_date"]]
    n_bad_logic = before - len(df)

    df["days_late"] = (df["payment_date"] - df["due_date"]).dt.days
    df.loc[df["cheque_bounced"] == True, "days_late"] = np.nan
    df["is_eid_ramzan_period"] = df["due_date"].apply(is_seasonal)

    final_cols = ["transaction_id", "dealer_id", "invoice_date", "due_date", "payment_date",
                  "amount_pkr", "payment_method", "cheque_bounced", "days_late", "is_eid_ramzan_period"]
    for c in final_cols:
        if c not in df.columns:
            df[c] = np.nan
    clean = df[final_cols].reset_index(drop=True)

    report.update({
        "output_rows": len(clean),
        "retention_rate": round(len(clean) / report["input_rows"], 3) if report["input_rows"] else 0,
        "dropped_missing_dealer": n_missing,
        "dropped_orphan_dealer": n_orphan,
        "dropped_bad_dates": n_bad_dates,
        "dropped_bad_amount": n_bad_amount,
        "dropped_bad_bounce_flag": n_bad_bool,
        "dropped_logical_inconsistency": n_bad_logic,
    })
    return clean, report


# ---------------------------------------------------------------------------
# FEATURE ENGINEERING + SCORING (ported from score_dealers.py)
# ---------------------------------------------------------------------------
def loo_group_rate(df, group_col, rate_col):
    s = df.groupby(group_col)[rate_col].transform("sum")
    c = df.groupby(group_col)[rate_col].transform("count")
    return ((s - df[rate_col]) / (c - 1).replace(0, np.nan)).fillna(df[rate_col].mean())


def compute_features(dealers: pd.DataFrame, salesmen: pd.DataFrame, txns: pd.DataFrame):
    """Returns (feat_df, insufficient_history_dealer_ids)."""
    dealers = dealers.copy()
    dealers["onboarding_date"] = pd.to_datetime(dealers["onboarding_date"])
    feature_window = txns[txns["due_date"] < CUTOFF_DATE]

    feat_rows, insufficient = [], []
    for dealer_id, dealer_row in dealers.set_index("dealer_id").iterrows():
        dgrp = feature_window[feature_window["dealer_id"] == dealer_id]
        if len(dgrp) < MIN_INVOICES:
            insufficient.append(dealer_id)
            continue
        non_bounced = dgrp[dgrp["cheque_bounced"] == False]
        bounce_rate_lifetime = dgrp["cheque_bounced"].mean()
        avg_days_late = non_bounced["days_late"].mean() if len(non_bounced) else 0
        nonseasonal = non_bounced[non_bounced.get("is_eid_ramzan_period", False) == False]
        avg_days_late_nonseasonal = nonseasonal["days_late"].mean() if len(nonseasonal) else avg_days_late
        payment_volatility = non_bounced["days_late"].std() if len(non_bounced) > 1 else 0
        years_active = max((CUTOFF_DATE.date() - dealer_row["onboarding_date"].date()).days / 365.25, 0)
        real_exposure_pkr = dealer_row["credit_limit_pkr"] * ((1 - ANNUAL_PKR_EROSION) ** years_active)
        mid = feature_window["invoice_date"].median()
        early = dgrp[dgrp["invoice_date"] < mid]
        late = dgrp[dgrp["invoice_date"] >= mid]
        order_frequency_trend = (len(late) - len(early)) / max(len(early), 1)
        feat_rows.append({
            "dealer_id": dealer_id, "dealer_name": dealer_row["dealer_name"],
            "city": dealer_row["city"], "sector": dealer_row["sector"],
            "salesman_id": dealer_row["salesman_id"], "territory_risk_tier": dealer_row["territory_risk_tier"],
            "is_salesman_favorite": bool(dealer_row["is_salesman_favorite"]),
            "credit_limit_pkr": dealer_row["credit_limit_pkr"],
            "bounce_rate_lifetime": bounce_rate_lifetime,
            "avg_days_late_nonseasonal": avg_days_late_nonseasonal,
            "payment_volatility": payment_volatility,
            "real_exposure_pkr": real_exposure_pkr,
            "order_frequency_trend": order_frequency_trend,
        })

    feat_df = pd.DataFrame(feat_rows)
    if len(feat_df) == 0:
        return feat_df, insufficient

    feat_df["salesman_default_rate_loo"] = loo_group_rate(feat_df, "salesman_id", "bounce_rate_lifetime")
    feat_df["territory_default_rate_loo"] = loo_group_rate(feat_df, "territory_risk_tier", "bounce_rate_lifetime")
    z_late = (feat_df["avg_days_late_nonseasonal"] - feat_df["avg_days_late_nonseasonal"].mean()) / feat_df["avg_days_late_nonseasonal"].std()
    z_vol = (feat_df["payment_volatility"] - feat_df["payment_volatility"].mean()) / feat_df["payment_volatility"].std()
    feat_df["payment_delay_severity"] = (z_late + z_vol) / 2
    feat_df = feat_df.merge(salesmen[["salesman_id", "salesman_name"]], on="salesman_id", how="left")

    return feat_df, insufficient


def score_with_model(feat_df: pd.DataFrame, artifact: dict) -> pd.DataFrame:
    """Applies the persisted model/scaler exactly as score_dealers.py does. No .fit() calls."""
    model, scaler, FEATURE_COLS = artifact["model"], artifact["scaler"], artifact["feature_columns"]
    score_params = artifact["score_params"]

    X_score_s = scaler.transform(feat_df[FEATURE_COLS])
    feat_df = feat_df.copy()
    feat_df["risk_probability"] = model.predict_proba(X_score_s)[:, 1]

    factor = score_params["pdo"] / np.log(2)
    offset = score_params["base_score"] - factor * np.log(score_params["base_odds"])
    def prob_to_score(p):
        p = np.clip(p, 1e-6, 1 - 1e-6)
        return np.clip(offset + factor * np.log((1 - p) / p), 300, 900)
    feat_df["credit_score"] = prob_to_score(feat_df["risk_probability"]).round(0)
    feat_df["risk_flag"] = pd.cut(feat_df["credit_score"], bins=[0, 580, 700, 900],
                                    labels=["RED", "AMBER", "GREEN"])

    contributions = X_score_s * model.coef_[0]
    contrib_df = pd.DataFrame(contributions, columns=FEATURE_COLS, index=feat_df.index)

    def top_reasons(idx, n=3):
        row = contrib_df.loc[idx]
        row_sorted = row.reindex(row.abs().sort_values(ascending=False).index)
        return [{"factor": REASON_LABELS[f], "direction": "increases_risk" if v > 0 else "reduces_risk", "weight": round(float(v), 3)}
                for f, v in row_sorted.head(n).items()]
    feat_df["top_reasons"] = [top_reasons(i) for i in feat_df.index]
    feat_df["scoring_method"] = "Statistical"
    return feat_df


def cold_start_score(dealers: pd.DataFrame, txns: pd.DataFrame, insufficient_ids: list) -> pd.DataFrame:
    """Rule-based provisional scoring for dealers with insufficient history."""
    if not insufficient_ids:
        return pd.DataFrame(columns=["dealer_id", "dealer_name", "city", "sector", "salesman_id", "salesman_name",
                                       "is_salesman_favorite", "credit_limit_pkr", "credit_score",
                                       "risk_flag", "risk_probability", "scoring_method", "top_reasons"])
    dealer_bounce = txns.groupby("dealer_id")["cheque_bounced"].mean()
    dwr = dealers.set_index("dealer_id").join(dealer_bounce.rename("bounce_rate"))
    dwr["bounce_rate"] = dwr["bounce_rate"].fillna(0)
    salesman_track = dwr.groupby("salesman_id")["bounce_rate"].mean()
    territory_base = dwr.groupby("territory_risk_tier")["bounce_rate"].mean()
    sector_base = dwr.groupby("sector")["bounce_rate"].mean()

    all_blended = []
    for _, d in dealers.iterrows():
        sr = salesman_track.get(d["salesman_id"], dealer_bounce.mean())
        tr = territory_base.get(d["territory_risk_tier"], dealer_bounce.mean())
        secr = sector_base.get(d["sector"], dealer_bounce.mean())
        all_blended.append(0.5 * sr + 0.3 * tr + 0.2 * secr)
    bmin, bmax = min(all_blended), max(all_blended)

    rows = []
    for dealer_id in insufficient_ids:
        d = dealers[dealers["dealer_id"] == dealer_id].iloc[0]
        sr = salesman_track.get(d["salesman_id"], dealer_bounce.mean())
        tr = territory_base.get(d["territory_risk_tier"], dealer_bounce.mean())
        secr = sector_base.get(d["sector"], dealer_bounce.mean())
        blended = 0.5 * sr + 0.3 * tr + 0.2 * secr
        norm = (blended - bmin) / (bmax - bmin + 1e-9)
        score = np.clip(680 - (norm * 200), 480, 680)
        tier = "AMBER-CAUTION" if score < 560 else "AMBER-STANDARD"
        rows.append({
            "dealer_id": dealer_id, "dealer_name": d["dealer_name"], "city": d["city"], "sector": d["sector"],
            "salesman_id": d["salesman_id"], "salesman_name": None, "is_salesman_favorite": bool(d["is_salesman_favorite"]),
            "credit_limit_pkr": d["credit_limit_pkr"], "credit_score": round(score), "risk_flag": tier,
            "risk_probability": None, "scoring_method": "Provisional (Cold-Start)",
            "top_reasons": [{"factor": "Insufficient payment history — score based on salesman/territory/sector averages", "direction": None, "weight": None}],
        })
    return pd.DataFrame(rows)

NAVY = RGBColor(0x1E, 0x27, 0x61)
RED = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xB8, 0x86, 0x0B)
GREEN = RGBColor(0x1E, 0x84, 0x49)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

TIER_COLOR = {"RED": RED, "AMBER": AMBER, "GREEN": GREEN}
TIER_LABEL = {"RED": "RED — High Risk", "AMBER": "AMBER — Moderate", "GREEN": "GREEN — Reliable"}
ACTION = {
    "RED": "Do not increase credit limit. Consider requiring partial upfront payment or cash-on-delivery terms on future orders.",
    "AMBER": "Maintain current credit limit. Monitor closely and reassess in 3 months.",
    "GREEN": "Eligible for credit limit review or increase based on strong payment history.",
}


def set_cell_background(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def fix_zoom_element(doc):
    """python-docx's default template ships a <w:zoom> element missing the
    required w:percent attribute -- present even in a totally blank document.
    Fixed here rather than left as a validation warning."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    # Same default-template issue: compatibilityMode is hardcoded to 14
    # (Word 2010), which triggers Word's "Compatibility Mode" banner on
    # every generated file. Bump to 15 (Word 2013+) to match modern Word.
    compat = settings.find(qn("w:compat"))
    if compat is not None:
        for cs in compat.findall(qn("w:compatSetting")):
            if cs.get(qn("w:name")) == "compatibilityMode":
                cs.set(qn("w:val"), "15")


def remove_table_borders(table):
    """Explicitly strips borders rather than relying on default table-style
    behavior, which can render visible gridlines depending on the Word/
    LibreOffice version -- guarantees the clean, borderless look regardless.
    OOXML's CT_TblPrBase schema requires child elements in a strict order;
    tblBorders must be inserted before shd/tblLayout/tblCellMar/tblLook if
    those already exist, not just appended at the end."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    following_tags = [qn(f"w:{t}") for t in
                       ("shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription")]
    insert_before = None
    for child in tbl_pr:
        if child.tag in following_tags:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(borders)
    else:
        tbl_pr.append(borders)

def tier_of(risk_flag: str) -> str:
    if risk_flag.startswith("RED"):
        return "RED"
    if risk_flag.startswith("GREEN"):
        return "GREEN"
    return "AMBER"
    
def build_risk_card_docx(dealer: dict) -> bytes:
    tier = tier_of(dealer["risk_flag"])
    color = TIER_COLOR[tier]
    is_contradiction = dealer.get("is_salesman_favorite") and tier == "RED"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    letterhead = doc.add_paragraph()
    r = letterhead.add_run("[ DISTRIBUTOR LETTERHEAD ]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    title = doc.add_paragraph()
    r = title.add_run("Dealer Credit Risk Card")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK

    sub = doc.add_paragraph()
    r = sub.add_run("Generated from historical payment data — internal decision support only")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    doc.add_paragraph()

    id_table = doc.add_table(rows=0, cols=2)
    remove_table_borders(id_table)
    id_rows = [
        ("Dealer Name", dealer["dealer_name"]),
        ("City / Territory", dealer["city"]),
        ("Sector", dealer["sector"]),
        ("Assigned Salesman", f"{dealer.get('salesman_name') or dealer['salesman_id']} ({dealer['salesman_id']})"),
    ]
    if is_contradiction:
        id_rows.append(("Salesman Relationship", "Marked as a trusted / favorite account"))
    for label, value in id_rows:
        row = id_table.add_row()
        row.cells[0].width = Inches(2.2)
        row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = GREY
        run = row.cells[1].paragraphs[0].add_run(str(value))
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RED if label == "Salesman Relationship" else DARK

    doc.add_paragraph()

    score_table = doc.add_table(rows=1, cols=2)
    score_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(score_table)
    left_cell, right_cell = score_table.rows[0].cells
    set_cell_background(left_cell, "F2F2F2")
    set_cell_background(right_cell, "F2F2F2")

    p = left_cell.paragraphs[0]
    r = p.add_run("CREDIT SCORE\n")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r2 = p.add_run(f"{round(dealer['credit_score'])}\n")
    r2.bold = True
    r2.font.size = Pt(34)
    r2.font.color.rgb = color
    r3 = p.add_run("out of 300–900")
    r3.font.size = Pt(8)
    r3.font.color.rgb = GREY

    p2 = right_cell.paragraphs[0]
    r4 = p2.add_run("RISK FLAG\n")
    r4.font.size = Pt(9)
    r4.font.color.rgb = GREY
    r5 = p2.add_run(f"{TIER_LABEL[tier]}\n")
    r5.bold = True
    r5.font.size = Pt(16)
    r5.font.color.rgb = color
    conf = dealer.get("risk_probability")
    conf_text = (f"Model confidence (ranking, not exact probability): {conf*100:.1f}%"
                 if conf is not None else
                 "Provisional score — based on portfolio averages, not this dealer's own history")
    r6 = p2.add_run(conf_text)
    r6.font.size = Pt(8)
    r6.font.color.rgb = GREY

    doc.add_paragraph()

    h = doc.add_paragraph()
    r = h.add_run("Key Contributing Factors")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    for i, reason in enumerate(dealer.get("top_reasons", []), 1):
        direction = reason.get("direction")
        prefix = "Increases risk" if direction == "increases_risk" else "Reduces risk" if direction == "reduces_risk" else ""
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {prefix + ' — ' if prefix else ''}{reason['factor']}")
        r.bold = True
        r.font.size = Pt(10.5)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r = h2.add_run("Recommended Action")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    action_text = ACTION[tier]
    if is_contradiction:
        action_text += (" Recommend a direct conversation with the assigned salesman given this "
                         "account's trusted status contradicts the payment record.")
    p = doc.add_paragraph()
    p.add_run(action_text).font.size = Pt(10.5)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    r = disclaimer.add_run(
        "This score is generated from historical payment behavior only and does not replace "
        "human judgment. Risk tier (RED/AMBER/GREEN) reflects the model's relative ranking "
        "within this portfolio; the exact numeric score will sharpen in precision as more "
        "historical data becomes available."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = GREY

    buf = io.BytesIO()
    fix_zoom_element(doc)
    doc.save(buf)
    return buf.getvalue()